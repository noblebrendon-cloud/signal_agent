import os
from openpyxl import Workbook
from docx import Document
from docx.shared import Pt
import fpdf

def generate_roi_calculator(base_dir):
    wb = Workbook()
    ws = wb.active
    ws.title = "Drift ROI Calculator"
    
    # Headers
    ws.append(["Signal Agent Drift ROI Calculator", "", ""])
    ws.append([])
    
    # Inputs
    ws.append(["INPUTS", "", ""])
    ws.append(["Queries Per Second (QPS)", 50, "Adjust this to match your normal load"])
    ws.append(["Silent Failure / Retry Rate (%)", 0.05, "5% fail/retry rate"])
    ws.append(["Cost Per Minute of Defective AI Output ($)", 250, "Industry avg context loss cost*"])
    ws.append([])
    
    # Formulas/Outputs
    # Calculate downtime cost. If QPS = 50, and 5% fail, that's 2.5 fails/sec = 150 fails/min.
    # But usually downtime cost is fixed per minute of active 'silent failure' state.
    ws.append(["ANNUAL METRICS", "", ""])
    ws.append(["Est. Defective Transactions per Year", "=B4*B5*60*60*24*365", ""])
    ws.append(["Estimated Annual Business Loss ($)", "=B6*B5*1000", "Total financial cost of drift execution"])
    ws.append([])
    
    ws.append(["SAVINGS", "", ""])
    ws.append(["Estimated Savings @ 20% Drift Instability Reduction ($)", "=B11*0.20", "Value added by early halt/circuit breaking"])
    
    # Citations
    ws.append([])
    ws.append(["*Footnote: Enterprise AI downtime and silent failure states cost an average of $250/min according to standard IT Ops SLA models for dynamic applications."])
    
    out_path = os.path.join(base_dir, "ROI_CALCULATOR.xlsx")
    wb.save(out_path)
    print("Saved Excel", out_path)

def generate_remediation(base_dir):
    doc = Document()
    doc.add_heading('Drift Diagnostic Remediation Outline', 0)
    
    p = doc.add_paragraph("Based on the captured Coherence Kernel metrics during the diagnostic, "
                          "we classify the required stability interventions into the following 3 core vectors:")
    
    doc.add_heading('Category 1: Retraining or Threshold Recalibration', level=1)
    doc.add_paragraph("TRIGGER CONDITION: If Φ1 (Distribution Divergence) > threshold (e.g., KL Divergence heavily shifts "
                      "without an immediately apparent hardware drop).")
    doc.add_paragraph("ACTION: The model's baseline semantic understanding is drifting away from reality. You must trigger "
                      "a delta-retrain on the novel edge cases or widen the confidence safety thresholds on the API layer "
                      "so that predictions flag uncertain inputs for human review.")

    doc.add_heading('Category 2: Hysteresis Tuning or Input Guardrails', level=1)
    doc.add_paragraph("TRIGGER CONDITION: If Φ2 (Performance Degradation) > threshold.")
    doc.add_paragraph("ACTION: The business output is already failing. Immediately implement input pre-filtering guardrails "
                      "(e.g., structural JSON validation, length limits, content filtering) to prevent toxic/malformed data "
                      "from reaching the endpoint, and tune the circuit breaker hysteresis to 'trip' faster under pressure.")

    doc.add_heading('Category 3: Immediate Circuit Breakdown (Failover)', level=1)
    doc.add_paragraph("TRIGGER CONDITION: If both Φ1 and Φ2 exceed their maximum tolerances forming an unstable feedback loop.")
    doc.add_paragraph("ACTION: Hard failure. Enforce structural routing to a fallback model (e.g., routing traffic to GPT-4o if a "
                      "Llama endpoint decays) while the primary node is quarantined and analyzed.")

    out_path = os.path.join(base_dir, "REMEDIATION_OUTLINE.docx")
    doc.save(out_path)
    print("Saved Docx", out_path)

class PDF(fpdf.FPDF):
    def footer(self):
        self.set_y(-15)
        self.set_font('helvetica', 'I', 8)
        self.cell(0, 10, 'Signal Agent Diagnostics', 0, 0, 'C')

def generate_security_pdf(base_dir):
    pdf = PDF()
    pdf.add_page()
    pdf.set_font("helvetica", "B", 16)
    pdf.cell(0, 10, "DATA SECURITY ONE PAGER", ln=True, align="C")
    pdf.ln(10)
    
    pdf.set_font("helvetica", "B", 12)
    pdf.cell(0, 8, "Strict Information Security Policy", ln=True)
    pdf.set_font("helvetica", "", 10)
    text = (
        "The $7K Drift Diagnostic is designed defensively to protect organizational intelligence. "
        "We do not exfiltrate your data, train on your weights, or retain your logs.\n\n"
        "1. On-Premise Execution: The diagnostic script runs strictly within your secure perimeter (AWS/GCP/Azure) "
        "via an isolated Docker container or raw static binary. No outbound data transmission occurs.\n\n"
        "2. No Data Retention: The audit strictly evaluates numeric probability distributions (logit maps and confidence floats). "
        "Raw inference text, PII, and output payload strings are never written to disk or preserved.\n\n"
        "3. TLS Required: If the diagnostic interfaces with an external vendor API endpoint, it strictly enforces TLS 1.3 "
        "with hardened cipher suites.\n\n"
        "4. Logs Purged Post-Delivery: Any temporary caching or metric aggregates used to compile the Drift Proof Document "
        "are systematically wiped using 3-pass erasure within 24 hours of final report delivery.\n\n"
        "Sample Mutual NDA Excerpt:\n"
        "\"Receiving Party (Signal Agent) agrees that all execution metrics, endpoint structures, and logic captured "
        "during the diagnostic engagement constitute strictly Confidential Information. Upon engagement closure, Receiving Party "
        "shall permanently destroy all logs, metrics, and environment variables derived from Disclosing Party's infrastructure.\""
    )
    pdf.multi_cell(0, 6, text)
    out_path = os.path.join(base_dir, "DATA_SECURITY_ONE_PAGER.pdf")
    pdf.output(out_path)
    print("Saved Security PDF", out_path)

def generate_offer_pdf(base_dir):
    pdf = PDF()
    pdf.add_page()
    pdf.set_font("helvetica", "B", 16)
    pdf.cell(0, 10, "OFFER SCOPE: $7K DRIFT DIAGNOSTIC", ln=True, align="C")
    pdf.ln(10)

    pdf.set_font("helvetica", "", 11)
    text = (
        "Engagement Format: 5-7 Day Async Target Audit\n"
        "Focus: Core System Drift Only (Phi 1 and Phi 2 Metrics)\n\n"
        "Our sole objective is to deterministically prove if your Tier-1 AI endpoint is silently bleeding accuracy "
        "due to data drift, prompt decay, or model degeneration.\n\n"
        "DELIVERABLES INCLUDED:\n"
        "- 1x Custom Distribution Divergence Graph (Phi 1 Baseline vs Production)\n"
        "- 1x Downstream Accuracy Correlation Analysis (Phi 2 Impact)\n"
        "- 1x Deterministic Remediation Brief (3 distinct architectural fixes mapping to your drift signature)\n"
        "- Full source output of `drift_metrics.csv` for internal usage\n\n"
        "FIXED ENGAGEMENT FEE:\n"
        "$7,000 USD (Flat Fee)\n\n"
        "PAYMENT TERMS:\n"
        "Invoiced Post-Delivery. You do not pay until the final Drift Proof report reaches your desk and the analysis is conclusive. "
        "This is an isolated diagnostic; there are no retainers or enforced upsells."
    )
    pdf.multi_cell(0, 6, text)
    out_path = os.path.join(base_dir, "OFFER_SCOPE.pdf")
    pdf.output(out_path)
    print("Saved Offer PDF", out_path)

if __name__ == "__main__":
    base = os.path.dirname(__file__)
    generate_roi_calculator(base)
    generate_remediation(base)
    generate_security_pdf(base)
    generate_offer_pdf(base)
    print("All artifacts generated.")
